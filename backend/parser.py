"""
parser.py — PDF and DOCX Text Extractor
Uses PyMuPDF (fitz) for PDFs and python-docx for DOCX files.
"""

import fitz  # PyMuPDF
import re
import io
from docx import Document


def _clean_text(raw_text: str) -> str:
    """
    Shared text cleanup: collapse blank lines and strip whitespace.
    """
    lines = [line.rstrip() for line in raw_text.splitlines()]
    cleaned_lines = []
    prev_blank = False
    for line in lines:
        is_blank = line.strip() == ""
        if is_blank and prev_blank:
            continue
        cleaned_lines.append(line)
        prev_blank = is_blank
    return "\n".join(cleaned_lines).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract and clean plain text from a PDF given its raw bytes.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        A clean plain-text string containing all readable text from the PDF.

    Raises:
        ValueError: If the PDF cannot be opened or contains no extractable text.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise ValueError(f"Failed to open PDF: {exc}") from exc

    pages_text = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pages_text.append(page.get_text("text"))

    doc.close()

    cleaned_text = _clean_text("\n".join(pages_text))

    if not cleaned_text:
        raise ValueError("No extractable text found in the PDF.")

    return cleaned_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract and clean plain text from a DOCX file given its raw bytes.

    Steps:
      1. Open the DOCX from an in-memory bytes buffer.
      2. Iterate over all paragraphs and extract their text.
      3. Also extract text from tables (row by row, cell by cell).
      4. Clean up excessive whitespace and blank lines.
      5. Return the cleaned plain-text string.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        A clean plain-text string containing all readable text from the DOCX.

    Raises:
        ValueError: If the DOCX cannot be opened or contains no extractable text.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f"Failed to open DOCX: {exc}") from exc

    text_parts = []

    # Step 2 — Extract text from all paragraphs
    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Step 3 — Extract text from all tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            text_parts.append(row_text)

    cleaned_text = _clean_text("\n".join(text_parts))

    if not cleaned_text:
        raise ValueError("No extractable text found in the DOCX.")

    return cleaned_text


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file type by extension and extract text accordingly.

    Supports: .pdf, .docx

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename: Original filename (used to detect extension).

    Returns:
        Extracted plain text string.

    Raises:
        ValueError: If the file type is unsupported or extraction fails.
    """
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{filename}'. Please upload a PDF or DOCX file.")
